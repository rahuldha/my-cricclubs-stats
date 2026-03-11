"""
Script to create MCP Server Guide Word Document
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    """Set background color for a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_code_block(doc, code_text):
    """Add a formatted code block."""
    para = doc.add_paragraph()
    run = para.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    # Add light gray background effect via border
    return para


def create_table(doc, headers, rows, header_color='2980B9'):
    """Create a formatted table with headers."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(header_cells[i], header_color)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = str(cell_data)

    doc.add_paragraph()  # Space after table
    return table


def main():
    doc = Document()

    # Title
    title = doc.add_heading('MCP Server Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Understanding the Model Context Protocol')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True

    doc.add_paragraph('Using the CricClubs Ground Stats MCP Server as an Example')
    doc.add_paragraph()

    # =========================================================================
    # Section 1: Simple Explanation
    # =========================================================================
    doc.add_heading('What is an MCP Server? (Simple Explanation)', level=1)

    doc.add_paragraph(
        'Imagine Claude is a really smart assistant who knows a lot of things. '
        'But Claude cannot leave its "room" - Claude can only talk to you.'
    )

    doc.add_heading('The Problem', level=2)
    doc.add_paragraph(
        'You ask Claude: "What are the cricket stats from CricClubs website?"\n'
        'Claude says: "I don\'t know, I can\'t visit websites on my own!"'
    )

    doc.add_heading('The Solution: MCP Server (the Helper)', level=2)
    doc.add_paragraph(
        'An MCP server is like a helper assistant who CAN go outside and fetch information. '
        'When you ask Claude for cricket stats, Claude asks the MCP server, '
        'the MCP server goes to CricClubs, gets the data, and brings it back to Claude.'
    )

    doc.add_paragraph('In short: An MCP server teaches Claude new tricks that Claude couldn\'t do alone.')

    # =========================================================================
    # Section 2: What is MCP
    # =========================================================================
    doc.add_heading('What is MCP?', level=1)

    doc.add_heading('Definition', level=2)
    doc.add_paragraph(
        'MCP stands for Model Context Protocol. It is a standard way for AI assistants '
        '(like Claude) to connect to external tools and data sources. Think of it like '
        'a universal plug/adapter that lets Claude talk to other programs.'
    )

    doc.add_heading('Who Created It?', level=2)
    doc.add_paragraph(
        'Anthropic (the company that makes Claude) created MCP and released it as an '
        'open standard in late 2024.'
    )

    doc.add_heading('Why Was It Created?', level=2)
    doc.add_paragraph(
        'Before MCP, every AI tool had its own custom way of connecting to external services. '
        'It was messy - like having different chargers for every phone. MCP provides one '
        'standard method that works for all tools.'
    )

    # =========================================================================
    # Section 3: The Core Idea
    # =========================================================================
    doc.add_heading('The Core Idea', level=1)

    doc.add_paragraph(
        'MCP defines a contract - a set of rules that both sides agree to follow. '
        'Claude Code (the client) says "I will send requests in THIS format" and '
        'the MCP Server (your tool) says "I will respond in THIS format". '
        'Both sides speak the same "language" - that\'s the standard.'
    )

    # =========================================================================
    # Section 4: The 3 Parts of an MCP Server
    # =========================================================================
    doc.add_heading('The 3 Parts of an MCP Server', level=1)

    # Part 1
    doc.add_heading('Part 1: Server Declaration', level=2)
    doc.add_paragraph('This tells Claude: "I\'m an MCP server with this name and purpose"')
    add_code_block(doc, '''from mcp.server.fastmcp import FastMCP

mcp = FastMCP(
    "CricClubs Ground Stats",                    # Server name
    instructions="Fetch match duration stats..." # What it does
)''')

    # Part 2
    doc.add_heading('Part 2: Tool Definition', level=2)
    doc.add_paragraph(
        'The @mcp.tool() decorator says: "This function is a tool Claude can use". '
        'The function name, parameters, types, and docstring all follow the MCP standard.'
    )
    add_code_block(doc, '''@mcp.tool()
def get_ground_stats(series_url: str) -> str:
    """
    Fetch match duration statistics aggregated by ground.

    Args:
        series_url: A CricClubs series URL containing league and clubId params.
                    Example: https://www.cricclubs.com/ARCL/listMatches.do?league=321&clubId=992
    """
    # ... your code here ...
    return "| Ground | Matches | Avg |..."''')

    doc.add_paragraph('The standard requires these elements:')
    create_table(doc,
        ['Element', 'What It Tells Claude', 'Our Example'],
        [
            ['Function name', 'What to call', 'get_ground_stats'],
            ['Parameters + types', 'What inputs are needed', 'series_url: str'],
            ['Docstring', 'How to use it', '"Fetch match duration..."'],
            ['Return type', 'What comes back', 'str (markdown table)'],
        ]
    )

    # Part 3
    doc.add_heading('Part 3: Transport (How They Communicate)', level=2)
    doc.add_paragraph(
        '"stdio" means communicate through standard input/output (like a terminal). '
        'This is how Claude Code and the MCP server exchange messages.'
    )
    add_code_block(doc, '''if __name__ == "__main__":
    mcp.run(transport="stdio")''')

    # =========================================================================
    # Section 5: What Makes It Standardized
    # =========================================================================
    doc.add_heading('What Makes It "Standardized"?', level=1)

    doc.add_paragraph(
        'The MCP standard defines exactly how messages look. All messages use JSON-RPC format. '
        'Here\'s what happens when you ask Claude for cricket stats:'
    )

    doc.add_heading('Step 1: Claude Sends a Request', level=2)
    add_code_block(doc, '''{
  "jsonrpc": "2.0",
  "id": 1,
  "method": "tools/call",
  "params": {
    "name": "get_ground_stats",
    "arguments": {
      "series_url": "https://www.cricclubs.com/ARCL/listMatches.do?league=321&clubId=992"
    }
  }
}''')

    doc.add_heading('Step 2: Your Server Responds', level=2)
    add_code_block(doc, '''{
  "jsonrpc": "2.0",
  "id": 1,
  "result": {
    "content": [
      {
        "type": "text",
        "text": "| Ground | Matches | Avg |\\n|--------|---------|-----|..."
      }
    ]
  }
}''')

    doc.add_paragraph(
        'This format is the standard. Every MCP server in the world uses this exact format. '
        'The mcp Python library handles all this formatting automatically - you just write '
        'your function and add @mcp.tool()!'
    )

    # =========================================================================
    # Section 6: Why Standardization Matters
    # =========================================================================
    doc.add_heading('Why Standardization Matters', level=1)

    doc.add_heading('Without a Standard', level=2)
    doc.add_paragraph(
        'Every server would have its own format. Claude would need custom code for each server:'
    )
    add_code_block(doc, '''Server A: {"tool": "get_stats", "input": "url"}
Server B: {"function": "fetch", "param": "link"}
Server C: {"action": "retrieve", "arg1": "address"}

Claude would need custom code for EACH server!''')

    doc.add_heading('With MCP Standard', level=2)
    doc.add_paragraph(
        'All servers use the same format. Claude uses one method for all servers:'
    )
    add_code_block(doc, '''Server A: {"method": "tools/call", "params": {"name": "...", "arguments": {...}}}
Server B: {"method": "tools/call", "params": {"name": "...", "arguments": {...}}}
Server C: {"method": "tools/call", "params": {"name": "...", "arguments": {...}}}

Claude uses ONE method for ALL servers!''')

    # =========================================================================
    # Section 7: What MCP Servers Can Do
    # =========================================================================
    doc.add_heading('What Can MCP Servers Do?', level=1)

    doc.add_paragraph('MCP servers can extend Claude\'s capabilities in many ways:')

    create_table(doc,
        ['MCP Server Type', 'What It Does'],
        [
            ['Database server', 'Let Claude query your database'],
            ['File server', 'Let Claude read/write files'],
            ['API server', 'Let Claude call web APIs'],
            ['CricClubs server (ours)', 'Let Claude fetch cricket statistics'],
            ['Email server', 'Let Claude send/read emails'],
            ['Calendar server', 'Let Claude manage calendar events'],
        ]
    )

    # =========================================================================
    # Section 8: Checklist
    # =========================================================================
    doc.add_heading('Checklist: Building a Standard MCP Server', level=1)

    doc.add_paragraph('To create a properly standardized MCP server, you need:')

    create_table(doc,
        ['Requirement', 'How To Do It', 'Example'],
        [
            ['1. Import MCP library', 'from mcp.server.fastmcp import FastMCP', 'Line 17 in our code'],
            ['2. Create server instance', 'mcp = FastMCP("Name")', 'Line 27 in our code'],
            ['3. Define tools with decorator', '@mcp.tool()', 'Line 160 in our code'],
            ['4. Type hints on parameters', 'def func(param: str)', 'series_url: str'],
            ['5. Docstring explaining usage', '"""Description..."""', 'Lines 161-167'],
            ['6. Return value', 'Return string/dict/list', 'Returns markdown'],
            ['7. Run with transport', 'mcp.run(transport="stdio")', 'Line 236 in our code'],
        ]
    )

    # =========================================================================
    # Section 9: Real-World Analogy
    # =========================================================================
    doc.add_heading('Real-World Analogy', level=1)

    doc.add_paragraph('MCP is like USB for AI:')
    doc.add_paragraph('• USB lets any device connect to any computer')
    doc.add_paragraph('• MCP lets any tool connect to any AI assistant')

    doc.add_paragraph(
        'This is why it matters: You build one MCP server, put it on GitHub, '
        'and anyone in the world can use it with Claude!'
    )

    # =========================================================================
    # Section 10: Summary
    # =========================================================================
    doc.add_heading('Summary', level=1)

    create_table(doc,
        ['Question', 'Answer'],
        [
            ['What is MCP?', 'A protocol (set of rules) for AI ↔ tools communication'],
            ['What\'s in an MCP server?', 'Server setup + tool definitions + transport'],
            ['Why "standardized"?', 'Everyone uses the same JSON message format'],
            ['What makes it standard?', 'Using the mcp library and following its patterns'],
            ['Who created it?', 'Anthropic (makers of Claude) in late 2024'],
        ]
    )

    # =========================================================================
    # Section 11: Our CricClubs Example
    # =========================================================================
    doc.add_heading('Our CricClubs MCP Server', level=1)

    doc.add_paragraph(
        'The CricClubs Ground Stats MCP server we built demonstrates all these concepts. '
        'When you ask Claude "Get ground stats for [CricClubs URL]", here\'s what happens:'
    )

    doc.add_paragraph('1. Claude receives your question')
    doc.add_paragraph('2. Claude recognizes it needs the get_ground_stats tool')
    doc.add_paragraph('3. Claude sends a standard JSON-RPC request to the MCP server')
    doc.add_paragraph('4. The MCP server fetches data from CricClubs website')
    doc.add_paragraph('5. The MCP server returns the results in standard JSON-RPC format')
    doc.add_paragraph('6. Claude displays the cricket statistics to you')

    doc.add_paragraph()
    doc.add_paragraph('The MCP standard makes all of this seamless and automatic!')

    # Save document
    output_path = 'MCP_Server_Guide.docx'
    doc.save(output_path)
    print(f'Word document saved to: {output_path}')


if __name__ == '__main__':
    main()
